import docx

doc = docx.Document()

doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS DE AGENCIAMENTO COMERCIAL E INTERMEDIAÇÃO DE NEGÓCIOS', level=1)

doc.add_heading('PREÂMBULO DO CONTRATO', level=2)
doc.add_paragraph('CONTRATANTE: ViDi EDUCAÇÃO LTDA, pessoa jurídica de direito privado, inscrita no CNPJ sob o nº 64.465.973/0001-89, sediada na Av. PL3, nº 960, Complemento Quadra H4, Lote 01/03, Sala 2508, Edifício T, Bairro Lozandes, Goiânia, Goiás, CEP 74.884-115, neste ato representada legalmente por seus administradores LUIZ EDUARDO ARAUJO PORTAL e CAMYLLA BORGES CONCEIÇÃO BOUFLEUR, doravante denominada simplesmente CONTRATANTE.')
doc.add_paragraph('CONTRATADA: [RAZÃO SOCIAL DA EMPRESA DA MARIANA], pessoa jurídica de direito privado, inscrita no CNPJ sob o nº [00.000.000/0000-00], estabelecida no endereço [ENDEREÇO COMPLETO DA MARIANA], neste ato representada por sua sócia titular MARIANA [SOBRENOME DA MARIANA], inscrita no CPF sob o nº [000.000.000-00].')
doc.add_paragraph('Resolvem, de comum acordo, firmar este contrato de prestação de serviços comerciais, que se regerá pelas cláusulas e condições seguintes:')

doc.add_heading('DO OBJETO (RESUMO)', level=2)
table1 = doc.add_table(rows=2, cols=3)
table1.style = 'Table Grid'
hdr_cells1 = table1.rows[0].cells
hdr_cells1[0].text = 'Item'
hdr_cells1[1].text = 'Entregável'
hdr_cells1[2].text = 'Descrição'
row_cells1 = table1.rows[1].cells
row_cells1[0].text = '1.0'
row_cells1[1].text = 'Agenciamento Comercial e Intermediação'
row_cells1[2].text = 'Atuação na prospecção ativa e passiva de clientes, qualificação de leads, apresentação de produtos/serviços, negociação e fechamento de contratos. Gestão de carteira e CRM.'

doc.add_heading('DA REMUNERAÇÃO (RESUMO)', level=2)
table2 = doc.add_table(rows=3, cols=3)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Item'
hdr_cells2[1].text = 'Descritivo'
hdr_cells2[2].text = 'Percentual / Valor'
row1_cells2 = table2.rows[1].cells
row1_cells2[0].text = '1.0'
row1_cells2[1].text = 'Comissão sobre Vendas'
row1_cells2[2].text = '10% (dez por cento) sobre o valor líquido recebido.'
row2_cells2 = table2.rows[2].cells
row2_cells2[0].text = '2.0'
row2_cells2[1].text = 'Ajuda de Custo / Fixos'
row2_cells2[2].text = 'Não aplicável (R$ 0,00).'

doc.add_heading('COMUNICAÇÕES', level=2)
table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Table Grid'
hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'Pela Contratante'
hdr_cells3[1].text = 'Pela Contratada'
row1_cells3 = table3.rows[1].cells
row1_cells3[0].text = 'Luiz Eduardo A. Portal / Camylla B. C. Boufleur'
row1_cells3[1].text = 'Nome: Mariana [Sobrenome da Mariana]'
row2_cells3 = table3.rows[2].cells
row2_cells3[0].text = 'E-mail: [INSERIR E-MAIL DA ViDi]'
row2_cells3[1].text = 'E-mail: mariana@vidiceo.com.br'
row3_cells3 = table3.rows[3].cells
row3_cells3[0].text = 'Telefone: [INSERIR TELEFONE DA ViDi]'
row3_cells3[1].text = 'Telefone: [TELEFONE DA MARIANA]'

doc.add_heading('CONDIÇÕES GERAIS DE CONTRATAÇÃO', level=2)
doc.add_heading('CLÁUSULA PRIMEIRA – DO OBJETO E ATUAÇÃO', level=3)
doc.add_paragraph('1.1. O objeto deste contrato é a prestação de serviços de Agenciamento Comercial e Intermediação de Negócios pela CONTRATADA à CONTRATANTE, visando a expansão da carteira de clientes da ViDi.')
doc.add_paragraph('1.2. A CONTRATADA atuará com autonomia técnica e administrativa, sem subordinação jurídica, executando prospecção, qualificação, negociação e atualização de CRM.')
doc.add_paragraph('1.3. Da Territorialidade: A CONTRATADA exercerá suas atividades em território nacional, sem exclusividade de zona.')
doc.add_paragraph('1.4. DA NÃO EXCLUSIVIDADE (Cláusula Expressa):')
doc.add_paragraph('A CONTRATADA não atuará com exclusividade para a CONTRATANTE. Fica expressamente autorizado à CONTRATADA prestar serviços a terceiros, inclusive a empresas do mesmo segmento ou concorrentes da CONTRATANTE, bem como representar outros produtos de alto ticket e parcerias as quais ela possua interesse, desde que, nessa atuação, não utilize informações confidenciais, segredos comerciais ou bases de dados de propriedade da CONTRATANTE, sob pena das sanções previstas na Cláusula 13.')

doc.add_heading('CLÁUSULA SEGUNDA - DO PRAZO', level=3)
doc.add_paragraph('2.1. O prazo de validade deste contrato é por tempo indeterminado, com início na data de sua assinatura.')

doc.add_heading('CLÁUSULA TERCEIRA - DA REMUNERAÇÃO E COMISSIONAMENTO', level=3)
doc.add_paragraph('3.1. A CONTRATADA será remunerada exclusivamente por êxito (Performance), inexistindo fixo ou ajuda de custo.')
doc.add_paragraph('3.2. Do Comissionamento: A CONTRATANTE pagará à CONTRATADA a comissão de 10% (dez por cento) sobre o valor líquido das vendas efetivadas.')
doc.add_paragraph('3.3. Condição de Pagamento: O pagamento é condicionado ao efetivo recebimento dos valores pela CONTRATANTE (Regime de Caixa). Em vendas parceladas, o pagamento será proporcional à liquidação das parcelas.')
doc.add_paragraph('3.4. Cancelamento/Inadimplência: Em caso de não pagamento pelo cliente final, o repasse da comissão será suspenso e eventuais adiantamentos serão estornados.')
doc.add_paragraph('3.5. Apuração mensal (dia 1 ao 30), Nota Fiscal até o dia 05 e pagamento até o dia 10.')

doc.add_heading('CLÁUSULA QUARTA - DAS OBRIGAÇÕES DA CONTRATADA', level=3)
doc.add_paragraph('4.1. Executar a prospecção e manter o CRM atualizado diariamente;')
doc.add_paragraph('4.2. Responder integralmente por seus sócios, prepostos e encargos trabalhistas de sua equipe;')
doc.add_paragraph('4.3. Respeitar o sigilo e não desviar clientela.')

doc.add_heading('CLÁUSULA QUINTA - DAS OBRIGAÇÕES DA CONTRATANTE', level=3)
doc.add_paragraph('5.1. Pagar as comissões devidas e fornecer acesso às ferramentas (CRM) e materiais.')

doc.add_heading('CLÁUSULA SEXTA - DO SIGILO E CONFIDENCIALIDADE', level=3)
doc.add_paragraph('6.1. Dever de sigilo absoluto sobre estratégias, leads e dados, vigente por prazo indeterminado.')

doc.add_heading('CLÁUSULA SÉTIMA – DA NÃO SOLICITAÇÃO E NÃO CONCORRÊNCIA', level=3)
doc.add_paragraph('7.1. Extinto o contrato, a CONTRATADA não poderá aliciar funcionários ou clientes da carteira ativa da CONTRATANTE pelo prazo de 02 (dois) anos.')

doc.add_heading('CLÁUSULA OITAVA – PROTEÇÃO DE DADOS (LGPD)', level=3)
doc.add_paragraph('8.1. Cumprimento integral da LGPD, usando dados de clientes apenas para as vendas deste contrato.')

doc.add_heading('CLÁUSULA NONA – PROPRIEDADE DA CARTEIRA', level=3)
doc.add_paragraph('9.1. Todo cliente ou lead cadastrado no CRM pertence à CONTRATANTE. É vedada a apropriação da base de dados pela CONTRATADA após a rescisão.')

doc.add_heading('CLÁUSULA DÉCIMA – COMPLIANCE', level=3)
doc.add_paragraph('10.1. Observância às leis anticorrupção e ética comercial.')

doc.add_heading('CLÁUSULA DÉCIMA PRIMEIRA - DA RESCISÃO', level=3)
doc.add_paragraph('11.1. Rescisão imotivada mediante aviso prévio de 30 (trinta) dias.')
doc.add_paragraph('11.2. Sem comissão de cauda (exceto vendas já fechadas e implantadas antes do desligamento).')
doc.add_paragraph('11.3. Rescisão por Justa Causa: Ocorrendo violação de sigilo, desvio de carteira ou retenção de valores, o contrato será rescindido imediatamente.')
doc.add_paragraph('11.3.1. Neste caso, a CONTRATANTE fica autorizada a reter comissões pendentes até o limite do valor da multa estipulada na Cláusula 13, para compensação de danos.')

doc.add_heading('CLÁUSULA DÉCIMA SEGUNDA – ELEIÇÃO DO FORO ARBITRAL', level=3)
doc.add_paragraph('12.1. Litígios resolvidos por arbitragem na 6ª Câmara de Conciliação, Mediação e Arbitragem de Goiânia (6ª CCMA/GO).')

doc.add_heading('CLÁUSULA DÉCIMA TERCEIRA – PENALIDADES (CLÁUSULA OBJETIVA)', level=3)
doc.add_paragraph('13.1. Em caso de violação das obrigações de Confidencialidade (Cláusula 6ª), Não Solicitação (Cláusula 7ª), Propriedade da Carteira (Cláusula 9ª) ou Uso Indevido de Dados, a Parte infratora pagará à outra multa compensatória equivalente à soma das comissões recebidas nos últimos 03 (três) meses de vigência do contrato.')
doc.add_paragraph('13.1.1. Caso o contrato tenha menos de 3 meses, ou a média seja inferior a R$ 5.000,00, fixa-se a multa mínima no valor de R$ 5.000,00 (cinco mil reais).')
doc.add_paragraph('13.1.2. O pagamento da multa não isenta a infratora de indenizar perdas e danos suplementares comprovados que excedam o valor da penalidade.')

doc.add_heading('CLÁUSULA DÉCIMA QUARTA - DISPOSIÇÕES GERAIS', level=3)
doc.add_paragraph('14.1. Inexistência de vínculo empregatício e natureza estritamente comercial.')

doc.add_paragraph('\nE, por estarem justas e contratadas, assinam as partes.\n')
doc.add_paragraph('[LOCAL], [DATA].\n')

doc.add_paragraph('_____________________________________________________')
doc.add_paragraph('ViDi EDUCAÇÃO LTDA\nContratante')

doc.add_paragraph('\n_____________________________________________________')
doc.add_paragraph('[RAZÃO SOCIAL DA EMPRESA DA MARIANA]\nContratada (repr. legal: Mariana [Sobrenome da Mariana])')

doc.add_paragraph('\nTESTEMUNHAS:\n')

doc.add_paragraph('1. _______________________________ \n   Nome: \n   CPF:')
doc.add_paragraph('2. _______________________________ \n   Nome: \n   CPF:')

doc.save("c:\\Users\\clebe\\OneDrive\\Área de Trabalho\\ProjetosGit\\TecnoIT\\Projeto Vidi\\Contrato_Mariana.docx")
